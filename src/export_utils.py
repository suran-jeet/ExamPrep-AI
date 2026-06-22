from __future__ import annotations

from io import BytesIO
import re


def as_markdown(title: str, body: str) -> str:
    clean_title = title.strip() or "ExamPrep AI Output"
    return f"# {clean_title}\n\n{body.strip()}\n"


def as_text(title: str, body: str) -> str:
    clean_title = title.strip() or "ExamPrep AI Output"
    underline = "=" * len(clean_title)
    text = re.sub(r"^#+\s*", "", body.strip(), flags=re.MULTILINE)
    text = text.replace("**", "")
    return f"{clean_title}\n{underline}\n\n{text}\n"


def as_pdf(title: str, body: str) -> bytes:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
    except Exception as exc:
        raise RuntimeError("Install reportlab to export PDF files.") from exc

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=42,
        leftMargin=42,
        topMargin=42,
        bottomMargin=42,
    )
    styles = getSampleStyleSheet()
    story = [Paragraph(_escape(title), styles["Title"]), Spacer(1, 12)]

    for block in body.strip().split("\n\n"):
        normalized = block.strip()
        if not normalized:
            continue
        if normalized.startswith("#"):
            heading = normalized.lstrip("#").strip()
            story.append(Paragraph(_escape(heading), styles["Heading2"]))
        else:
            lines = normalized.splitlines()
            html = "<br/>".join(_escape(line) for line in lines)
            story.append(Paragraph(html, styles["BodyText"]))
        story.append(Spacer(1, 8))

    document.build(story)
    return buffer.getvalue()


def as_docx(title: str, body: str) -> bytes:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("Install python-docx to export DOCX files.") from exc

    document = Document()
    document.add_heading(title.strip() or "ExamPrep AI Output", level=0)

    for block in body.strip().split("\n\n"):
        normalized = block.strip()
        if not normalized:
            continue
        if normalized.startswith("#"):
            heading = normalized.lstrip("#").strip()
            level = min(4, max(1, len(normalized) - len(normalized.lstrip("#"))))
            document.add_heading(heading, level=level)
            continue
        for line in normalized.splitlines():
            clean = line.strip()
            if clean.startswith("- "):
                document.add_paragraph(clean[2:], style="List Bullet")
            else:
                document.add_paragraph(clean.replace("**", ""))

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _escape(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )
